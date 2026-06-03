#!/usr/bin/env python3
"""
Generate reference.docx for pandoc with coursework-compliant styles.

Starts from pandoc's built-in reference.docx (ensures Table/Compact styles
exist so tables do not collapse in Word), then patches typography and layout.

Usage (from repo root):
    REPO_ROOT=<path> OUT=<path> python3 build/scripts/make_reference_docx.py
"""
import os
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Mm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(os.environ.get("REPO_ROOT", Path(__file__).parent.parent.parent))
OUT  = Path(os.environ.get("OUT",       ROOT / "build" / "reference.docx"))

# ── Typography constants ─────────────────────────────────────────────────────
FONT      = "Times New Roman Cyr"
BODY_PT   = 14
TABLE_PT  = 12

# Geometry helpers (all Word measurements in twips: 1 in = 1440 twips)
_MM = 1440 / 25.4

def _mm(v: float) -> int:  return int(round(v * _MM))
def _pt(v: float) -> int:  return int(v * 20)   # twentieths-of-a-point (spacing)
def _hp(v: float) -> int:  return int(v * 2)    # half-points (font sz)

# Page A4 + margins
PAGE_W = _mm(210)   # 11906
PAGE_H = _mm(297)   # 16838
MAR_L  = _mm(25)    # 1417
MAR_R  = _mm(15)    #  850
MAR_T  = _mm(20)    # 1134
MAR_B  = _mm(20)    # 1134

# Paragraph indent 1.25 cm
INDENT = _mm(1.25)  # 709

# Line spacing
LINE15 = 360  # 1.5× (auto)
LINE10 = 240  # 1.0× (single)

# ── Low-level XML helpers ────────────────────────────────────────────────────

def _find(el, tag: str):
    return el.find(qn(tag))


def _ensure(el, tag: str):
    child = _find(el, tag)
    if child is None:
        child = OxmlElement(tag)
        el.append(child)
    return child


def _set(el, tag: str, **attrs):
    """Ensure child tag exists and set all w:-namespaced attributes."""
    child = _ensure(el, tag)
    for k, v in attrs.items():
        child.set(qn(k), str(v))
    return child


def _clear(el, tag: str):
    for c in list(el.findall(qn(tag))):
        el.remove(c)


def _pPr(style_el):
    pPr = _find(style_el, "w:pPr")
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        style_el.insert(0, pPr)
    return pPr


def _rPr(style_el):
    rPr = _find(style_el, "w:rPr")
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style_el.append(rPr)
    return rPr


# ── Style property setters ───────────────────────────────────────────────────

def set_font(style_el, size_pt: int, bold=False, caps=False):
    rPr = _rPr(style_el)
    _set(rPr, "w:rFonts", **{
        "w:ascii": FONT, "w:hAnsi": FONT,
        "w:cs": FONT,    "w:eastAsia": FONT,
    })
    _set(rPr, "w:sz",   **{"w:val": str(_hp(size_pt))})
    _set(rPr, "w:szCs", **{"w:val": str(_hp(size_pt))})
    if bold:
        _ensure(rPr, "w:b")
        _ensure(rPr, "w:bCs")
    else:
        _clear(rPr, "w:b")
        _clear(rPr, "w:bCs")
    if caps:
        _ensure(rPr, "w:caps")
    else:
        _clear(rPr, "w:caps")
    # Remove formatting that should not be in body text
    for tag in ("w:i", "w:iCs", "w:u", "w:strike", "w:color", "w:highlight"):
        _clear(rPr, tag)


def set_para(style_el,
             align="both",
             first_line=INDENT,
             line=LINE15,
             before=0,
             after=0,
             page_break=False,
             keep_next=False):
    pPr = _pPr(style_el)

    _set(pPr, "w:jc", **{"w:val": align})

    ind = _ensure(pPr, "w:ind")
    ind.set(qn("w:left"), "0")
    ind.attrib.pop(qn("w:hanging"), None)
    ind.attrib.pop(qn("w:right"), None)
    if first_line and first_line > 0:
        ind.set(qn("w:firstLine"), str(first_line))
    else:
        ind.set(qn("w:firstLine"), "0")

    sp = _ensure(pPr, "w:spacing")
    sp.set(qn("w:before"),    str(before))
    sp.set(qn("w:after"),     str(after))
    sp.set(qn("w:line"),      str(line))
    sp.set(qn("w:lineRule"),  "auto")
    sp.attrib.pop(qn("w:beforeAutospacing"), None)
    sp.attrib.pop(qn("w:afterAutospacing"),  None)

    if page_break:
        _ensure(pPr, "w:pageBreakBefore")
    else:
        _clear(pPr, "w:pageBreakBefore")

    if keep_next:
        _ensure(pPr, "w:keepNext")
    else:
        _clear(pPr, "w:keepNext")

    _clear(pPr, "w:contextualSpacing")


# ── Style accessor ───────────────────────────────────────────────────────────

def _style(doc, name: str):
    """Return existing style or add new paragraph style."""
    for s in doc.styles:
        if s.name == name:
            return s
    return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


# ── Per-style patches ────────────────────────────────────────────────────────

def patch_normal(doc):
    s = _style(doc, "Normal")
    set_font(s.element, BODY_PT)
    set_para(s.element, align="both", first_line=INDENT,
             line=LINE15, before=0, after=0)


def patch_heading1(doc):
    """Chapter headings: UPPERCASE, centered, bold, page break before."""
    s = _style(doc, "Heading 1")
    set_font(s.element, BODY_PT, bold=True, caps=True)
    set_para(s.element, align="center", first_line=0,
             line=LINE15, before=0, after=_pt(12),
             page_break=True, keep_next=True)


def patch_heading2(doc):
    """Subsection headings: bold, page break before, 36/24 pt spacing."""
    s = _style(doc, "Heading 2")
    set_font(s.element, BODY_PT, bold=True)
    set_para(s.element, align="both", first_line=INDENT,
             line=LINE15, before=_pt(36), after=_pt(24),
             page_break=True, keep_next=True)


def patch_heading3(doc):
    """Point headings: bold, no page break, 24/12 pt spacing."""
    s = _style(doc, "Heading 3")
    set_font(s.element, BODY_PT, bold=True)
    set_para(s.element, align="both", first_line=INDENT,
             line=LINE15, before=_pt(24), after=_pt(12),
             page_break=False, keep_next=True)


def patch_compact(doc):
    """
    'Compact' is the style pandoc uses for table-cell paragraphs.
    Without it table cells collapse to one column in Word/LibreOffice.
    12 pt, single spacing, no first-line indent.
    """
    s = _style(doc, "Compact")
    set_font(s.element, TABLE_PT)
    set_para(s.element, align="both", first_line=0,
             line=LINE10, before=0, after=0)


def patch_table_paragraph(doc):
    """Alternate table-cell style used in some pandoc versions."""
    s = _style(doc, "Table Paragraph")
    set_font(s.element, TABLE_PT)
    set_para(s.element, align="both", first_line=0,
             line=LINE10, before=0, after=0)


def patch_caption(doc):
    """
    'Caption' style: used via {custom-style="Caption"} for figure captions.
    Centered, no first-line indent, 14 pt, blank line above and below.
    """
    s = _style(doc, "Caption")
    set_font(s.element, BODY_PT)
    set_para(s.element, align="center", first_line=0,
             line=LINE15, before=_pt(6), after=_pt(6))


def add_custom_styles(doc):
    """
    Styles referenced via {custom-style="..."} in preprocessed markdown.

    Figure        – image container: centered, no indent, blank line before.
    TableCaption  – table caption above table: no indent, blank line before.
    """
    fig = _style(doc, "Figure")
    set_font(fig.element, BODY_PT)
    set_para(fig.element, align="center", first_line=0,
             line=LINE15, before=_pt(6), after=0)

    tc = _style(doc, "TableCaption")
    set_font(tc.element, BODY_PT)
    set_para(tc.element, align="both", first_line=0,
             line=LINE15, before=_pt(6), after=0)


# ── Page geometry ────────────────────────────────────────────────────────────

def patch_page_setup(doc):
    for section in doc.sections:
        section.page_width     = Twips(PAGE_W)
        section.page_height    = Twips(PAGE_H)
        section.left_margin    = Twips(MAR_L)
        section.right_margin   = Twips(MAR_R)
        section.top_margin     = Twips(MAR_T)
        section.bottom_margin  = Twips(MAR_B)
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)


# ── Page number header ───────────────────────────────────────────────────────

def add_page_number_header(doc):
    """Right-aligned PAGE field in the header of every section."""
    for section in doc.sections:
        header = section.header

        # Clear existing header content
        for para in header.paragraphs:
            para.clear()

        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = para.add_run()
        r = run._r

        def fld_char(kind: str):
            e = OxmlElement("w:fldChar")
            e.set(qn("w:fldCharType"), kind)
            return e

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "

        r.append(fld_char("begin"))
        r.append(instr)
        r.append(fld_char("separate"))
        r.append(fld_char("end"))

        # Apply font to the run
        rPr = OxmlElement("w:rPr")
        _set(rPr, "w:rFonts", **{"w:ascii": FONT, "w:hAnsi": FONT, "w:cs": FONT})
        _set(rPr, "w:sz",   **{"w:val": str(_hp(BODY_PT))})
        _set(rPr, "w:szCs", **{"w:val": str(_hp(BODY_PT))})
        r.insert(0, rPr)


# ── Entry point ──────────────────────────────────────────────────────────────

def main():
    print("Fetching pandoc built-in reference.docx...")
    proc = subprocess.run(
        ["pandoc", "--print-default-data-file", "reference.docx"],
        capture_output=True, check=True,
    )
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        f.write(proc.stdout)
        tmp = Path(f.name)

    doc = Document(str(tmp))
    tmp.unlink(missing_ok=True)

    print("Setting page geometry (A4, margins 25/15/20/20 mm)...")
    patch_page_setup(doc)

    print("Patching paragraph styles...")
    patch_normal(doc)
    patch_heading1(doc)
    patch_heading2(doc)
    patch_heading3(doc)
    patch_compact(doc)
    patch_table_paragraph(doc)
    patch_caption(doc)
    add_custom_styles(doc)

    print("Adding page-number header...")
    add_page_number_header(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
